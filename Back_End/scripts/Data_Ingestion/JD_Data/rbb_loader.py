import re
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from docx import Document

from jd_config import (
    DIVISION_RBB,
    RESP_STOP_KEYWORDS,
    RESP_START_KEYWORDS,
    RBB_JD_PATH,
)

# ─────────────────────────────────────────────────────────────
# Text cleaning helper
# ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Normalize extracted Word text."""

    replacements = {
        "â€™": "'",
        "â€œ": '"',
        "â€\x9d": '"',
        "â€“": "-",
        "â€”": "-",
        "Â": "",
        "ﬁ": "fi",
        "ﬂ": "fl",
    }

    for bad, good in replacements.items():
        text = text.replace(bad, good)

    text = re.sub(r"[\u2022\u25cf\u25a0\uf0b7]", "", text)
    text = re.sub(r"[^\w\s,.;:/&()'\-]", "", text)
    text = re.sub(r"\s+", " ", text)

    return text.strip()


# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────

def _parse_grade(raw: Optional[str]) -> Optional[int]:
    if not raw:
        return None
    m = re.search(r"(\d+)", raw)
    return int(m.group(1)) if m else None


def _is_stop(text: str) -> bool:
    lo = text.lower()
    return any(kw.lower() in lo for kw in RESP_STOP_KEYWORDS)


def _is_resp_header(text: str) -> bool:
    lo = text.lower()
    return any(kw.lower() in lo for kw in RESP_START_KEYWORDS)


# ─────────────────────────────────────────────────────────────
# Per-table parser
# ─────────────────────────────────────────────────────────────

def _parse_table(table) -> Optional[Dict]:
    field_paras: Dict[str, str] = {}
    resp_cell = None

    seen_cells = set()

    for row in table.rows:
        for cell in row.cells:
            cid = id(cell)
            if cid in seen_cells:
                continue
            seen_cells.add(cid)

            paras = [_clean_text(p.text) for p in cell.paragraphs if p.text.strip()]
            if not paras:
                continue

            first = paras[0]

            if _is_resp_header(first):
                resp_cell = cell
                continue

            for p in paras:
                if ":" in p:
                    label, _, val = p.partition(":")
                    field_paras[label.strip().lower()] = _clean_text(val)

    def get(labels):
        for label in labels:
            if label in field_paras and field_paras[label]:
                return field_paras[label]
        return ""

    job_title = get(["job title"])
    division = get(["division"]) or DIVISION_RBB
    unit = get([ "unit"])
    department = get(["department"])
    job_grade = _parse_grade(get(["job grade"]))
    job_category = get(["job category"])
    job_objective = get(["job objective"])


    if not job_title:
        return None

    # ───────────────── responsibilities ─────────────────
    responsibilities: List[str] = []

    if resp_cell is not None:
        capturing = False

        for para in resp_cell.paragraphs:
            text = _clean_text(para.text)

            if not text:
                continue

            if _is_resp_header(text):
                capturing = True
                continue

            if not capturing:
                continue

            if _is_stop(text):
                break

            clean = re.sub(r"^[\-\*\•\uf0b7\s]+", "", text).rstrip(";").strip()

            if len(clean) > 8:
                responsibilities.append(clean)
    return {
        "source": "JD",
        "division": division,
        "unit": unit,
        "department": department,
        "job_title": job_title,
        "job_grade": job_grade,
        "job_category": job_category,
        "job_objective": job_objective,
        "responsibilities": responsibilities,
        "num_responsibilities": len(responsibilities),
    }


# ─────────────────────────────────────────────────────────────
# Loader class
# ─────────────────────────────────────────────────────────────

class RBBLoader:
    """Parse Retail & Branch Banking JD Word document."""

    def __init__(self, file_path: Optional[Path] = None):
        self.file_path = file_path or RBB_JD_PATH
        self._df: Optional[pd.DataFrame] = None
        print(f"📁 RBBLoader → {self.file_path.name}")

    def load(self) -> pd.DataFrame:
        print(f"\n📂 Parsing {self.file_path.name} ...")

        doc = Document(self.file_path)
        records = []

        for table in doc.tables:
            rec = _parse_table(table)

            if rec:
                records.append(rec)

                grade = str(rec["job_grade"]) if rec["job_grade"] is not None else "?"
                print(
                    f"✓ [{grade}] {rec['job_title']} "
                    f"({rec['num_responsibilities']} responsibilities)"
                )

        self._df = pd.DataFrame(records)

        print(
            f"\n✅ RBB: {len(self._df)} JDs | "
            f"{self._df['num_responsibilities'].sum()} total responsibilities"
        )

        return self._df

    def get_dataframe(self) -> pd.DataFrame:
        if self._df is None:
            self.load()
        return self._df